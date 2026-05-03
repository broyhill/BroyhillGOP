#!/usr/bin/env python3
"""
Build GOD_FILE_INDEX_V8.html
- Deep-indexes session .md content (headings, key phrases, SQL refs, dollar amounts)
- Auto-classifies files into ecosystems or suggests new ecosystem assignments
- Fixes OPEN button to launch correct app per file type
- Updates relay health check to new server 37.27.169.232
"""

import json, re, os
from pathlib import Path
from datetime import datetime
from collections import defaultdict
# ── Ecosystem classifier keywords ─────────────────────────────────────
ECO_CLASSIFIERS = {
    "E00": ["datahub","master database","data warehouse","etl","data pipeline","schema migration"],
    "E01": ["donor intelligence","donor scoring","donor grading","donor profile","cultivation","archetype","propensity"],
    "E02": ["donation processing","payment","winred","actblue","contribution","receipt","refund"],
    "E03": ["candidate profile","marketing automation","candidate page","campaign marketing"],
    "E04": ["activist","grassroots","mobilization","petition","rally"],
    "E05": ["volunteer management","volunteer tracking","volunteer hours","volunteer signup"],
    "E06": ["analytics engine","reporting","metrics","kpi","heat map","precinct analysis"],
    "E07": ["issue tracking","issue position","policy position","stance"],
    "E08": ["communications library","template","message template","comm library"],
    "E09": ["content creation","ai content","copywriting","blog post","press release"],
    "E10": ["compliance","fec compliance","campaign finance","disclosure","filing deadline"],
    "E11": ["budget management","training","lms","learning management","course","spending"],
    "E12": ["campaign operations","field operations","gotv","get out the vote","canvass"],
    "E13": ["ai hub","machine learning","model training","inference","prompt"],
    "E14": ["print production","print shop","mailer design","postcard"],
    "E15": ["contact directory","address book","contact lookup","phone list"],
    "E16": ["tv","radio","voice synthesis","text to speech","audio","broadcast ad"],
    "E17": ["rvm","ringless voicemail","voicemail drop"],
    "E18": ["vdp","variable data","print advertising","personalized mail"],
    "E19": ["social media","facebook","twitter","instagram","tiktok","linkedin","social post"],
    "E20": ["intelligence brain","brain","orchestrator","decision engine","reasoning"],
    "E21": ["ml clustering","kmeans","cluster analysis","segmentation model","dbscan"],
    "E22": ["a/b test","thompson sampling","experiment","variant","split test"],
    "E23": ["creative asset","3d","blender","render","graphic design","image generation"],
    "E24": ["candidate portal","candidate dashboard","candidate login"],
    "E25": ["donor portal","donor dashboard","donor login","my donations"],
    "E26": ["volunteer portal","volunteer dashboard","volunteer login"],
    "E27": ["realtime dashboard","live dashboard","websocket","real-time"],
    "E28": ["financial dashboard","revenue","burn rate","cash flow","financial report"],
    "E29": ["analytics dashboard","chart","visualization","graph","report viewer"],
    "E30": ["email engine","sendgrid","mailgun","email delivery","email template","newsletter"],
    "E31": ["sms engine","twilio","mms","text message","10dlc","short code"],
    "E32": ["phone banking","predictive dialer","call center","phone bank","auto dial"],
    "E33": ["direct mail","usps","postal","mail merge","bulk mail","lob"],
    "E34": ["events","event registration","rsvp","ticket","event page"],
    "E35": ["interactive comm","auto response","chatbot","conversational"],
    "E36": ["messenger","whatsapp","telegram","signal","chat integration"],
    "E37": ["event management","event planning","venue","catering","schedule event"],
    "E38": ["volunteer coordination","shift","assignment","volunteer dispatch"],
    "E39": ["p2p fundraising","peer to peer","fundraising page","crowdfund"],
    "E40": ["automation control","workflow automation","trigger","rule engine","cron"],
    "E41": ["campaign builder","campaign wizard","campaign setup","launch campaign"],
    "E42": ["news intelligence","news monitoring","media monitoring","press clip","opposition research"],
    "E43": ["advocacy","petition","grassroots advocacy","action alert","constituent"],
    "E44": ["vendor compliance","social intelligence","vendor management","security audit"],
    "E45": ["video studio","video editing","video production","video render","youtube"],
    "E46": ["broadcast hub","multi-channel","broadcast","mass communication"],
    "E47": ["ai script","unified voice","script generator","talking points","speech"],
    "E48": ["communication dna","messaging framework","tone analysis","brand voice"],
    "E49": ["interview system","gpu orchestrator","runpod","gpu scheduling"],
    "E50": ["gpu orchestrator","runpod","gpu cluster","inference server"],
    "E51": ["nexus dashboard","nexus","master dashboard","command center"],
    "E52": ["contact intelligence","contact enrichment","data enrichment","clearbit","geocodio"],
    "E53": ["document generation","pdf generation","docx generation","letter merge","certificate"],
    "E54": ["calendar","scheduling","appointment","booking","availability"],
    "E55": ["api gateway","api management","rate limit","authentication","oauth","webhook"],
    "E56": ["visitor deanonymization","website visitor","ip lookup","anonymous visitor"],
    "E57": ["messaging center","inbox","unified messaging","notification center"],
    "E58": ["business model","pricing","subscription","candidate network","data acquisition","saas"],
    # === BUG FIX (2026-05-03): add E59/E60/E61 to classifier ===
    # E59 corrected 2026-05-03 per Nexus drift log: E59 is Microsegment Intelligence,
    # NOT Nervous Net. Nervous Net keywords removed entirely; they need a home that
    # doesn't exist yet (deferred until blueprint audit reaches that question).
    "E59": ["microsegment","micro-segment","segment selection","segmentation",
            "archetype calibration","voter archetype","donor archetype",
            "intent classification","audience intelligence",
            "psychographic","demographic cluster",
            "evangelical","maga","establishment","libertarian","rural_ag",
            "defense_leo","anti_woke","healthcare_prof",
            "salesforce einstein"],
    "E60": ["poll","survey","county intensity","catawba","meredith","high point university","cygnal","civitas","yougov","voter archetype","onsp","issue intensity","poll source","nc county","calibration"],
    "E61": ["source ingestion","identity resolution","master data","mdm","source registry","entity resolution","raw_sources"],
    # === END FIX ===
}
ECOSYSTEM_NAMES = {
    "E00":"DataHub / Master Database","E01":"Donor Intelligence Engine","E02":"Donation Processing",
    "E03":"Candidate Profiles / Marketing","E04":"Activist Network","E05":"Volunteer Management",
    "E06":"Analytics Engine","E07":"Issue Tracking","E08":"Communications Library",
    "E09":"Content Creation AI","E10":"Compliance Manager","E11":"Budget / Training LMS",
    "E12":"Campaign Operations","E13":"AI Hub","E14":"Print Production",
    "E15":"Contact Directory","E16":"TV/Radio/Voice","E17":"RVM",
    "E18":"VDP / Print Ads","E19":"Social Media Manager","E20":"Intelligence Brain",
    "E21":"ML Clustering","E22":"A/B Testing","E23":"Creative Asset / 3D",
    "E24":"Candidate Portal","E25":"Donor Portal","E26":"Volunteer Portal",
    "E27":"Realtime Dashboard","E28":"Financial Dashboard","E29":"Analytics Dashboard",
    "E30":"Email Engine","E31":"SMS/MMS Engine","E32":"Phone Banking",
    "E33":"Direct Mail","E34":"Events","E35":"Interactive Comm Hub",
    "E36":"Messenger Integration","E37":"Event Management","E38":"Volunteer Coordination",
    "E39":"P2P Fundraising","E40":"Automation Control","E41":"Campaign Builder",
    "E42":"News Intelligence","E43":"Advocacy Tools","E44":"Vendor/Social Intel",
    "E45":"Video Studio","E46":"Broadcast Hub","E47":"AI Script/Voice",
    "E48":"Communication DNA","E49":"Interview System","E50":"GPU Orchestrator",
    "E51":"Nexus Dashboard","E52":"Contact Intelligence","E53":"Document Generation",
    "E54":"Calendar/Scheduling","E55":"API Gateway","E56":"Visitor Deanonymization",
    "E57":"Messaging Center","E58":"Business Model / Network",
    # === BUG FIX (2026-05-03): names for E59/E60/E61 ===
    "E59":"Microsegment Intelligence","E60":"Poll & Survey Intelligence","E61":"Source Ingestion / Identity",
    # === END FIX ===
}

# ── Read V7 data ──────────────────────────────────────────────────────
V7_PATH = "/Users/Broyhill/Documents/GitHub/BroyhillGOP/GOD_FILE_INDEX_V7.html"
with open(V7_PATH) as f:
    lines = f.readlines()
data_line = lines[174]
start = data_line.index('[')
end = data_line.rindex(']') + 1
FILES = json.loads(data_line[start:end])
print(f"Loaded {len(FILES)} files from V7")
# ── Deep content + ecosystem classifier ────────────────────────────────
def classify_ecosystem(text, filename):
    """Match content to known ecosystems, return list of matched eco IDs."""
    matched = []
    text_lower = (text + " " + filename).lower()
    for eco_id, keywords in ECO_CLASSIFIERS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score >= 2:  # need 2+ keyword hits
            matched.append(eco_id)
    return matched

def extract_deep_content(filepath):
    """Extract headings, dollar amounts, table names, key phrases from file."""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
    except:
        return "", []
    snippets = []
    headings = re.findall(r'^#{1,4}\s+(.+)$', text, re.MULTILINE)
    if headings:
        snippets.append("HEADINGS: " + " | ".join(h.strip() for h in headings[:30]))
    dollars = re.findall(r'\$[\d,]+(?:\.\d+)?[KMB]?', text)
    if dollars:
        snippets.append("AMOUNTS: " + " ".join(set(dollars[:20])))
    tables = re.findall(r'(?:public|core|staging|archive|norm|audit|raw)\.\w+', text)
    if tables:
        snippets.append("TABLES: " + " ".join(set(tables[:30])))
    names = re.findall(r'\b(?:BROYHILL|POPE|ESHELMAN|FAISON|TILLIS|GUSTAFSON|BOLIEK|GREEN|BURR|BUDD|MEADOWS)\b', text, re.I)
    if names:
        snippets.append("NAMES: " + " ".join(set(n.upper() for n in names)))
    tech_pats = [
        r'(?:person_spine|contribution_map|donor_cluster|contact_spine)',
        r'(?:PITR|TRUNCATE|ALTER TABLE|CREATE TABLE|INSERT INTO|identity.resolution)',
        r'(?:dedup|7.pass|canary|golden.record|DataTrust|Acxiom|RNC_RegID)',
        r'(?:relay|Hetzner|Supabase|Redis|Docker|migration)',
        r'(?:NCBOE|FEC|WinRed|NCGOP|microsegment|SIC|NAICS|honorific)',
        r'(?:Phase [A-K]|Stage [1-7]|Pass [1-8]|VoterFile|voter.file)',
    ]
    tech = []
    for pat in tech_pats:
        tech.extend(re.findall(pat, text, re.I))
    if tech:
        snippets.append("TECH: " + " ".join(set(t.lower() for t in tech[:40])))
    row_counts = re.findall(r'[\d,]+\s*(?:rows?|records?|active|inactive)', text, re.I)
    if row_counts:
        snippets.append("COUNTS: " + " | ".join(set(row_counts[:15])))
    eco_refs = re.findall(r'\bE\d{2}\b', text)
    if eco_refs:
        snippets.append("ECOSYSTEMS: " + " ".join(sorted(set(eco_refs))))
    eco_class = classify_ecosystem(text, os.path.basename(filepath))
    return " || ".join(snippets), eco_class

# === BUG FIX (2026-05-03): docx topic-extraction helper ===
def extract_docx_topics(filepath):
    """Parse .docx via python-docx and apply the same topic extraction used for .md/.sql/.py.
    Adds .docx files to topic-indexing without modifying extract_deep_content."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        return "", []
    try:
        doc = Document(filepath)
    except Exception:
        return "", []
    paras = [pp.text for pp in doc.paragraphs[:600] if pp.text and pp.text.strip()]
    for tbl in doc.tables[:5]:
        for row in tbl.rows[:50]:
            cells = [c.text.strip() for c in row.cells[:6] if c.text and c.text.strip()]
            if cells: paras.append(' | '.join(cells))
    text = '\n'.join(paras)
    if not text:
        return "", []
    snippets = []
    # Use first ~10 paragraphs as headings (DOCX rarely has markdown #)
    if paras:
        snippets.append('HEADINGS: ' + ' | '.join(p[:80] for p in paras[:10]))
    eco_refs = re.findall(r'\bE\d{2}\b', text)
    if eco_refs: snippets.append('ECOSYSTEMS: ' + ' '.join(sorted(set(eco_refs))))
    tables = re.findall(r'(?:public|core|staging|archive|norm|audit|raw|nc_)\.?\w+', text)
    if tables: snippets.append('TABLES: ' + ' '.join(sorted(set(tables))[:30]))
    dollars = re.findall(r'\$[\d,]+(?:\.\d+)?[KMB]?', text)
    if dollars: snippets.append('AMOUNTS: ' + ' '.join(sorted(set(dollars))[:20]))
    eco_class = classify_ecosystem(text, os.path.basename(filepath))
    return " || ".join(snippets), eco_class
# === END FIX ===

# === BUG FIX (2026-05-03): title extraction (H1 / first heading / filename fallback) ===
def extract_title(filepath, ftype):
    """Extract a clean title for display.
    Returns empty string if nothing distinct -- caller falls back to filename."""
    if not os.path.isfile(filepath):
        return ""
    try:
        if ftype == 'docx':
            try:
                from docx import Document
                doc = Document(filepath)
                for pp in doc.paragraphs[:60]:
                    if pp.style and 'Heading' in (pp.style.name or '') and pp.text.strip():
                        return pp.text.strip()[:140]
                for pp in doc.paragraphs[:60]:
                    if pp.text.strip():
                        return pp.text.strip()[:140]
                return ""
            except Exception:
                return ""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as fh:
            head = fh.read(8000)
        if ftype == 'md':
            m = re.search(r'^#\s+(.+?)$', head, re.MULTILINE)
            if m: return m.group(1).strip()[:140]
        if ftype == 'html':
            m = re.search(r'<title[^>]*>(.+?)</title>', head, re.IGNORECASE | re.DOTALL)
            if m: return m.group(1).strip()[:140]
            m = re.search(r'<h1[^>]*>(.+?)</h1>', head, re.IGNORECASE | re.DOTALL)
            if m: return re.sub(r'<[^>]+>', '', m.group(1)).strip()[:140]
        if ftype == 'py':
            m = re.search(r'^"""\s*\n?(.+?)\n', head, re.MULTILINE)
            if m: return m.group(1).strip()[:140]
        if ftype == 'sql':
            m = re.search(r'^--\s*(.+?)$', head, re.MULTILINE)
            if m: return m.group(1).strip()[:140]
        return ""
    except Exception:
        return ""
# === END FIX ===

# ── Human-readable category names ─────────────────────────────────────
def human_category(filepath, filetype, filename):
    """Return a plain-English category an amateur would understand."""
    p = filepath.lower()
    fn = filename.lower()
    # Code files
    if filetype == 'py':
        if 'ecosystem' in fn: return 'Python App (Ecosystem)'
        if 'pipeline' in fn or 'import' in fn: return 'Python Data Loader'
        if 'engine' in fn: return 'Python Engine'
        if 'test' in fn: return 'Python Test'
        if 'integration' in fn or 'api' in fn: return 'Python API Connector'
        return 'Python Script'
    if filetype == 'sql':
        if 'migration' in fn or 'alter' in fn: return 'Database Migration'
        if 'schema' in fn or 'create' in fn: return 'Database Schema'
        if 'fix' in fn or 'patch' in fn: return 'Database Fix'
        if 'audit' in fn or 'check' in fn: return 'Database Audit'
        return 'SQL Script'
    if filetype == 'html':
        if 'ecosystem' in fn: return 'App Screen (Ecosystem)'
        if 'portal' in fn or 'dashboard' in fn: return 'Dashboard Page'
        if 'index' in fn: return 'Homepage'
        return 'Web Page'
    if filetype == 'md':
        if 'session' in fn: return 'Session Transcript'
        if 'readme' in fn: return 'Instructions'
        if 'todo' in fn or 'task' in fn: return 'Task List'
        if 'plan' in fn or 'architecture' in fn: return 'Architecture Plan'
        if 'dedup' in fn or 'pipeline' in fn: return 'Pipeline Design'
        if 'audit' in fn or 'report' in fn: return 'Audit Report'
        if 'state' in fn: return 'Project State'
        return 'Documentation'
    if filetype in ('json',):
        if 'search' in fn or 'index' in fn: return 'Search Index Data'
        if 'config' in fn or 'setting' in fn: return 'Configuration'
        return 'JSON Data'
    if filetype in ('csv', 'xlsx', 'xls', 'numbers'):
        if 'donor' in fn or 'donation' in fn: return 'Donor Data File'
        if 'voter' in fn or 'datatrust' in fn: return 'Voter Data File'
        if 'fec' in fn: return 'FEC Contribution Data'
        if 'ncboe' in fn or 'boe' in fn: return 'NC BOE Donation Data'
        if 'schema' in fn or 'column' in fn: return 'Data Dictionary'
        return 'Spreadsheet / Data'
    if filetype == 'docx':
        if 'blueprint' in fn or 'ecosystem' in fn: return 'Ecosystem Blueprint'
        if 'report' in fn or 'audit' in fn: return 'Report Document'
        return 'Word Document'
    if filetype == 'sh': return 'Shell Script'
    if filetype == 'txt': return 'Text File'
    if filetype in ('js', 'jsx', 'ts'): return 'JavaScript Code'
    if filetype == 'css': return 'Stylesheet'
    return 'Other File'

# ── Human-readable file description ───────────────────────────────────
def describe_file(filename, filetype, eco_ids):
    """One-line plain-English description of what this file IS."""
    fn = filename.lower()
    desc_parts = []
    # Ecosystem match
    if eco_ids:
        names = [ECOSYSTEM_NAMES.get(e, e) for e in eco_ids[:3]]
        desc_parts.append("Related to: " + ", ".join(names))
    # File-specific clues
    if filetype == 'py':
        if 'import' in fn: desc_parts.append("Loads data into the database")
        elif 'engine' in fn: desc_parts.append("Background processing engine")
        elif 'search' in fn: desc_parts.append("Search/indexing tool")
        elif 'pipeline' in fn: desc_parts.append("Multi-step data pipeline")
        elif 'audit' in fn: desc_parts.append("Checks data quality")
    elif filetype == 'sql':
        if 'schema' in fn: desc_parts.append("Creates database tables")
        elif 'migration' in fn: desc_parts.append("Changes database structure")
        elif 'fix' in fn: desc_parts.append("Repairs database issue")
    elif filetype == 'html':
        if 'ecosystem' in fn: desc_parts.append("Interactive app screen")
        elif 'portal' in fn: desc_parts.append("User-facing dashboard")
    elif filetype == 'md':
        if 'session' in fn: desc_parts.append("AI work session record")
        elif 'plan' in fn: desc_parts.append("Implementation plan")
        elif 'dedup' in fn: desc_parts.append("Donor matching strategy")
    return " | ".join(desc_parts) if desc_parts else ""
# ── Enhance FILES ─────────────────────────────────────────────────────
enhanced = 0
new_eco_assignments = 0
for f in FILES:
    path = f.get('p', '')
    ftype = f.get('t', '')
    fname = f.get('n', '')
    # Human-readable category
    f['c'] = human_category(path, ftype, fname)
    # === BUG FIX (2026-05-03): extract title for display ===
    if ftype in ('md','html','py','sql','docx'):
        try:
            f['title'] = extract_title(path, ftype)
        except Exception:
            f['title'] = ''
    # === END FIX ===
    # Deep-index .md and .sql files
    deep, eco_class = "", []
    if ftype in ('md', 'sql', 'py') and os.path.isfile(path):
        deep, eco_class = extract_deep_content(path)
    # === BUG FIX (2026-05-03): .docx topic-extraction (parallel to existing) ===
    elif ftype == 'docx' and os.path.isfile(path):
        deep, eco_class = extract_docx_topics(path)
    # === END FIX ===
    if deep:
        existing_q = f.get('q', '')
        f['q'] = (existing_q + " || " + deep) if existing_q else deep
        f['dc'] = 1
        enhanced += 1
    # Auto-classify into ecosystems if not already tagged
    if eco_class:
        existing_ecos = set(f.get('e', '').split(',')) if f.get('e') else set()
        new_ecos = set(eco_class) - existing_ecos - {''}
        if new_ecos:
            all_ecos = sorted((existing_ecos | new_ecos) - {''})
            f['e'] = ','.join(all_ecos)
            new_eco_assignments += 1
    # File description
    eco_ids = f.get('e', '').split(',') if f.get('e') else []
    f['desc'] = describe_file(fname, ftype, eco_ids)

print(f"Deep-indexed: {enhanced} files")
print(f"New ecosystem assignments: {new_eco_assignments} files")
# ── Generate V8 HTML ──────────────────────────────────────────────────
FILES_JSON = json.dumps(FILES, separators=(',',':'))
ECO_NAMES_JS = json.dumps(ECOSYSTEM_NAMES)

OUTPUT = "/Users/Broyhill/Documents/GitHub/BroyhillGOP/GOD_FILE_INDEX_V8.html"

with open(OUTPUT, 'w', encoding='utf-8') as out:
    out.write('''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>BroyhillGOP GOD FILE V8</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{background:#0d1117;color:#e6edf3;font-family:'Segoe UI',system-ui,sans-serif;height:100vh;display:flex;flex-direction:column;overflow:hidden}
#topbar{background:#161b22;border-bottom:1px solid #30363d;padding:10px 16px;display:flex;align-items:center;gap:12px;flex-shrink:0;z-index:10}
#topbar h1{font-size:16px;font-weight:700;color:#f0883e;white-space:nowrap}
.relay-badge{font-size:11px;padding:2px 8px;border-radius:10px;background:#21262d;border:1px solid #30363d;color:#8b949e}
.relay-badge.ok{border-color:#3fb950;color:#3fb950}
#search-box{flex:1;background:#21262d;border:1px solid #30363d;border-radius:6px;padding:7px 12px;color:#e6edf3;font-size:14px;outline:none}
#search-box:focus{border-color:#388bfd}
#stats-bar{font-size:12px;color:#8b949e;white-space:nowrap}
#main{display:flex;flex:1;overflow:hidden}
#sidebar{width:270px;flex-shrink:0;background:#161b22;border-right:1px solid #30363d;overflow-y:auto;padding:12px}
#results{flex:1;overflow-y:auto;padding:12px}
.filter-section{margin-bottom:16px}
.filter-label{font-size:11px;font-weight:600;text-transform:uppercase;color:#8b949e;letter-spacing:.05em;margin-bottom:6px;display:flex;align-items:center;justify-content:space-between;cursor:pointer}
.filter-label .arrow{transition:transform .2s}
.filter-label.collapsed .arrow{transform:rotate(-90deg)}
.filter-body{display:block}.filter-body.hidden{display:none}
.pill-row{display:flex;flex-wrap:wrap;gap:4px}
.pill{font-size:11px;padding:3px 8px;border-radius:10px;background:#21262d;border:1px solid #30363d;cursor:pointer;color:#8b949e;transition:all .15s}
.pill:hover,.pill.active{background:#388bfd22;border-color:#388bfd;color:#79c0ff}
.pill.active{font-weight:600}
input[type=date]{background:#21262d;border:1px solid #30363d;color:#e6edf3;padding:4px 8px;border-radius:5px;font-size:12px;width:100%}
.date-row{display:flex;gap:6px;align-items:center;font-size:11px;color:#8b949e}
.eco-grid{display:grid;grid-template-columns:1fr 1fr;gap:3px;max-height:200px;overflow-y:auto}
.eco-chip{font-size:10px;padding:2px 5px;border-radius:4px;background:#21262d;border:1px solid #30363d;cursor:pointer;color:#8b949e;text-align:center;transition:all .15s}
.eco-chip:hover,.eco-chip.active{background:#f0883e22;border-color:#f0883e;color:#f0883e}
.toggle-row{display:flex;flex-direction:column;gap:4px}
.toggle-btn{display:flex;align-items:center;gap:8px;font-size:12px;color:#8b949e;cursor:pointer;padding:4px 6px;border-radius:5px;border:1px solid transparent}
.toggle-btn:hover{background:#21262d;border-color:#30363d}
.toggle-btn.active{color:#3fb950;border-color:#3fb950;background:#3fb95011}
.toggle-btn .dot{width:10px;height:10px;border-radius:50%;background:#30363d;flex-shrink:0}
.toggle-btn.active .dot{background:#3fb950}
.clear-btn{font-size:11px;color:#8b949e;background:none;border:1px solid #30363d;padding:2px 8px;border-radius:5px;cursor:pointer;margin-top:8px;width:100%}
.clear-btn:hover{border-color:#f85149;color:#f85149}
.card{background:#161b22;border:1px solid #30363d;border-radius:8px;padding:12px;margin-bottom:8px;transition:border-color .15s}
.card:hover{border-color:#388bfd44}
.card-header{display:flex;align-items:flex-start;gap:8px;margin-bottom:4px}
.file-name{font-size:13px;font-weight:600;color:#e6edf3;flex:1;word-break:break-all;line-height:1.3}
.type-badge{font-size:10px;padding:1px 6px;border-radius:4px;font-weight:700;flex-shrink:0;font-family:monospace}
.date-badge{font-size:10px;color:#8b949e;flex-shrink:0}
.card-desc{font-size:11px;color:#58a6ff;margin-bottom:4px;font-style:italic}
.card-meta{display:flex;align-items:center;gap:8px;margin-bottom:6px;flex-wrap:wrap}
.meta-item{font-size:10px;color:#8b949e}
.badge{font-size:9px;padding:1px 5px;border-radius:3px;font-weight:700;text-transform:uppercase}
.badge-session{background:#6e40c911;border:1px solid #6e40c9;color:#a371f7}
.badge-eco{background:#f0883e11;border:1px solid #f0883e;color:#f0883e}
.badge-deep{background:#3fb95011;border:1px solid #3fb950;color:#3fb950}
.badge-code{background:#388bfd11;border:1px solid #388bfd;color:#79c0ff}
.eco-tags{display:flex;flex-wrap:wrap;gap:3px;margin-bottom:6px}
.eco-tag{font-size:9px;padding:1px 5px;border-radius:3px;background:#f0883e22;border:1px solid #f0883e55;color:#f0883e;cursor:pointer}
.topic-tags{display:flex;flex-wrap:wrap;gap:3px;margin-bottom:8px}
.topic-chip{font-size:10px;padding:1px 6px;border-radius:8px;background:#388bfd11;border:1px solid #388bfd44;color:#79c0ff;cursor:pointer}
.deep-snippet{font-size:10px;color:#8b949e;background:#0d1117;border:1px solid #21262d;border-radius:4px;padding:6px 8px;margin-bottom:8px;max-height:50px;overflow:hidden;cursor:pointer;transition:max-height .3s}
.deep-snippet.expanded{max-height:400px;overflow-y:auto}
.actions{display:flex;gap:5px;flex-wrap:wrap}
.btn{font-size:10px;padding:3px 9px;border-radius:4px;cursor:pointer;border:1px solid;text-decoration:none;display:inline-block;transition:all .15s;font-weight:500}
.btn-open{background:#21262d;border-color:#30363d;color:#8b949e}
.btn-open:hover{border-color:#3fb950;color:#3fb950}
.btn-reveal{background:#21262d;border-color:#30363d;color:#8b949e}
.btn-reveal:hover{border-color:#f0883e;color:#f0883e}
.btn-copy{background:#21262d;border-color:#30363d;color:#8b949e}
.btn-copy:hover{border-color:#388bfd;color:#79c0ff}
.btn-drive{background:#21262d;border-color:#30363d;color:#8b949e}
.btn-drive:hover{border-color:#34d399;color:#34d399}
.btn.copied{border-color:#3fb950!important;color:#3fb950!important}
#load-more{width:100%;padding:10px;background:#21262d;border:1px solid #30363d;color:#8b949e;border-radius:6px;cursor:pointer;margin-top:8px}
#load-more:hover{border-color:#388bfd;color:#79c0ff}
.highlight{background:#f0883e44;border-radius:2px}
.no-results{text-align:center;color:#8b949e;padding:60px 20px}
::-webkit-scrollbar{width:6px}::-webkit-scrollbar-track{background:#0d1117}::-webkit-scrollbar-thumb{background:#30363d;border-radius:3px}
</style>
</head>
<body>
<div id="topbar">
  <h1>BROYHILLGOP GOD FILE V8</h1>
  <span class="relay-badge" id="relay-badge">RELAY</span>
  <span class="relay-badge" id="bridge-badge">BRIDGE</span>
  <input id="search-box" type="text" placeholder="Search files, topics, ecosystems, dollar amounts, table names..." autocomplete="off">
  <span id="stats-bar">Loading...</span>
</div>
<div id="main">
  <div id="sidebar">
    <div class="filter-section"><div class="filter-label" onclick="toggleSection(this)">CATEGORY <span class="arrow">&#9660;</span></div><div class="filter-body"><div class="pill-row" id="cat-pills"></div></div></div>
    <div class="filter-section"><div class="filter-label" onclick="toggleSection(this)">FILE TYPE <span class="arrow">&#9660;</span></div><div class="filter-body"><div class="pill-row" id="type-pills"></div></div></div>
    <div class="filter-section"><div class="filter-label" onclick="toggleSection(this)">DATE RANGE <span class="arrow">&#9660;</span></div><div class="filter-body"><div class="date-row"><span>From</span><input type="date" id="date-from"></div><div class="date-row" style="margin-top:4px"><span>To</span><input type="date" id="date-to"></div></div></div>
    <div class="filter-section"><div class="filter-label" onclick="toggleSection(this)">ECOSYSTEMS <span class="arrow">&#9660;</span></div><div class="filter-body"><div class="eco-grid" id="eco-grid"></div></div></div>
    <div class="filter-section"><div class="filter-label">QUICK FILTERS</div><div class="toggle-row">
      <div class="toggle-btn" id="toggle-session" onclick="toggleFilter('session')"><span class="dot"></span>Session Transcripts</div>
      <div class="toggle-btn" id="toggle-ecoref" onclick="toggleFilter('ecoref')"><span class="dot"></span>Ecosystem Files</div>
      <div class="toggle-btn" id="toggle-deep" onclick="toggleFilter('deep')"><span class="dot"></span>Deep-Indexed</div>
      <div class="toggle-btn" id="toggle-code" onclick="toggleFilter('code')"><span class="dot"></span>Code Files Only</div>
    </div></div>
    <button class="clear-btn" onclick="clearAllFilters()">CLEAR ALL FILTERS</button>
  </div>
  <div id="results"></div>
</div>
<script>
''')
    out.write(f'const FILES={FILES_JSON};\n')
    out.write(f'const ECO_NAMES={ECO_NAMES_JS};\n')
    out.write('''
const TYPE_COLORS={md:'#f0883e',py:'#3fb950',sql:'#a371f7',html:'#f778ba',json:'#79c0ff',csv:'#d2a8ff',docx:'#388bfd',xlsx:'#3fb950',txt:'#8b949e',sh:'#ffa657',js:'#f0e68c',ts:'#388bfd',jsx:'#f778ba',numbers:'#34d399'};
const APP_MAP={md:'Visual Studio Code',py:'Visual Studio Code',sql:'Visual Studio Code',js:'Visual Studio Code',ts:'Visual Studio Code',jsx:'Visual Studio Code',json:'Visual Studio Code',sh:'Visual Studio Code',css:'Visual Studio Code',html:'Google Chrome',htm:'Google Chrome',csv:'Numbers',xlsx:'Microsoft Excel',xls:'Microsoft Excel',numbers:'Numbers',docx:'Microsoft Word',doc:'Microsoft Word',pdf:'Preview',txt:'TextEdit',png:'Preview',jpg:'Preview'};
const PER_PAGE=40;
const state={search:'',cats:new Set(),types:new Set(),dateFrom:'',dateTo:'',ecos:new Set(),session:false,ecoref:false,deep:false,code:false,page:0,filtered:[]};

function buildSidebar(){
  const cats={},types={};
  FILES.forEach(f=>{cats[f.c]=(cats[f.c]||0)+1;types[f.t]=(types[f.t]||0)+1;});
  const catBox=document.getElementById('cat-pills');
  catBox.innerHTML=makePill('All','all-cat','cat',true);
  Object.entries(cats).sort((a,b)=>b[1]-a[1]).forEach(([c,n])=>{catBox.innerHTML+=makePill(c+' ('+n+')',c,'cat');});
  const typeBox=document.getElementById('type-pills');
  typeBox.innerHTML=makePill('All','all-type','type',true);
  Object.entries(types).sort((a,b)=>b[1]-a[1]).forEach(([t,n])=>{typeBox.innerHTML+=makePill('.'+t+' ('+n+')',t,'type');});
  const ecoBox=document.getElementById('eco-grid');
  for(let i=0;i<=58;i++){const id='E'+String(i).padStart(2,'0');ecoBox.innerHTML+=`<div class="eco-chip" data-eco="${id}" onclick="toggleEco('${id}')" title="${ECO_NAMES[id]||id}">${id}</div>`;}
}
function makePill(label,val,group,active=false){return `<span class="pill${active?' active':''}" data-val="${val}" data-pill-group="${group}" onclick="togglePill(this,'${group}')">${label}</span>`;}
// === ADDITION (2026-05-03): rollup counts -- dynamic per-pill match counts ===
function updatePillCounts(){
  const filtered=state.filtered;
  // Count matches per category and type
  const catCounts={},typeCounts={};
  filtered.forEach(f=>{catCounts[f.c]=(catCounts[f.c]||0)+1;typeCounts[f.t]=(typeCounts[f.t]||0)+1;});
  document.querySelectorAll('.pill').forEach(p=>{
    const grp=p.getAttribute('data-pill-group');
    const val=p.getAttribute('data-val');
    if(!grp||val==='all-'+grp)return;
    const n=(grp==='cat'?catCounts[val]:typeCounts[val])||0;
    // Strip any prior " (N)" suffix and re-append
    let txt=p.textContent.replace(/\s*\(\d[\d,]*\)\s*$/,'');
    p.textContent=txt+' ('+n.toLocaleString()+')';
  });
  // Eco chip counts
  const ecoCounts={};
  filtered.forEach(f=>{if(f.e){f.e.split(',').forEach(e=>{if(e)ecoCounts[e]=(ecoCounts[e]||0)+1;});}});
  document.querySelectorAll('.eco-chip').forEach(c=>{
    const id=c.getAttribute('data-eco');
    const n=ecoCounts[id]||0;
    c.setAttribute('title',(ECO_NAMES[id]||id)+' -- '+n+' matching');
    // Add a small superscript count
    let span=c.querySelector('.eco-cnt');
    if(!span){span=document.createElement('span');span.className='eco-cnt';span.style.cssText='font-size:8px;color:#484f58;margin-left:3px';c.appendChild(span);}
    span.textContent=n>0?'·'+n:'';
  });
}
// === END ADDITION ===

function applyFilters(){
  const q=state.search.toLowerCase().trim();
  const terms=q.split(/\\s+/).filter(Boolean);
  state.filtered=FILES.filter(f=>{
    if(state.cats.size&&!state.cats.has(f.c))return false;
    if(state.types.size&&!state.types.has(f.t))return false;
    if(state.session&&!f.x)return false;
    if(state.ecoref&&!f.e)return false;
    if(state.deep&&!f.dc)return false;
    if(state.code&&!['py','sql','js','ts','jsx','html','sh','css'].includes(f.t))return false;
    if(state.dateFrom&&f.d<state.dateFrom)return false;
    if(state.dateTo&&f.d>state.dateTo)return false;
    if(state.ecos.size){if(!f.e)return false;const fe=f.e.split(',');let m=false;state.ecos.forEach(e=>{if(fe.includes(e))m=true;});if(!m)return false;}
    if(terms.length){const hay=(f.n+' '+f.c+' '+(f.q||'')+' '+(f.e||'')+' '+(f.desc||'')).toLowerCase();return terms.every(t=>hay.includes(t));}
    return true;
  });
  state.page=0;
  document.getElementById('stats-bar').textContent=state.filtered.length+' / '+FILES.length+' files';
  // === ADDITION (2026-05-03): refresh rollup counts ===
  try{updatePillCounts();}catch(e){}
  // === END ADDITION ===
  render();
}

function render(){
  const box=document.getElementById('results');
  const end=(state.page+1)*PER_PAGE;
  const slice=state.filtered.slice(0,end);
  if(!slice.length){box.innerHTML='<div class="no-results">No files match your search.</div>';return;}
  const q=state.search.trim();
  box.innerHTML=slice.map(f=>renderCard(f,q)).join('');
  if(end<state.filtered.length)box.innerHTML+=`<button id="load-more" onclick="loadMore()">Load more (${state.filtered.length-end} remaining)</button>`;
}
function loadMore(){state.page++;render();}

function renderCard(f,q){
  const tc=TYPE_COLORS[f.t]||'#8b949e';
  const sz=f.s>1048576?(f.s/1048576).toFixed(1)+'MB':f.s>1024?(f.s/1024).toFixed(0)+'KB':f.s+'B';
  // === BUG FIX (2026-05-03): show extracted title in filename slot when available ===
  let name=esc(f.title&&f.title.length?f.title:f.n);
  // === END FIX ===
  if(q){q.split(' ').filter(Boolean).forEach(term=>{name=name.replace(new RegExp(escRe(term),'gi'),m=>`<span class="highlight">${m}</span>`);});}
  // Topics (original only, before ||)
  let topicHtml='';
  if(f.q){const raw=f.q.split('||')[0];const topics=raw.split(',').map(t=>t.trim()).filter(Boolean).slice(0,8);if(topics.length)topicHtml=`<div class="topic-tags">${topics.map(t=>`<span class="topic-chip" onclick="addTopicSearch('${esc(t)}')">${esc(t)}</span>`).join('')}</div>`;}
  // Deep snippet
  let deepHtml='';
  if(f.dc&&f.q&&f.q.includes('||')){let snippet=esc(f.q.split('||').slice(1).join(' ').trim()).substring(0,500);if(q)q.split(' ').filter(Boolean).forEach(term=>{snippet=snippet.replace(new RegExp(escRe(term),'gi'),m=>`<span class="highlight">${m}</span>`);});deepHtml=`<div class="deep-snippet" onclick="this.classList.toggle('expanded')" title="Click to expand">${snippet}</div>`;}
  // Ecosystems
  let ecoHtml='';
  if(f.e){const ecos=f.e.split(',').filter(Boolean);ecoHtml=`<div class="eco-tags">${ecos.map(e=>`<span class="eco-tag" title="${ECO_NAMES[e]||e}" onclick="filterByEco('${e}')">${e}</span>`).join('')}</div>`;}
  // Badges
  let badges='';
  if(f.x)badges+=`<span class="badge badge-session">TRANSCRIPT</span> `;
  if(f.e)badges+=`<span class="badge badge-eco">ECOSYSTEM</span> `;
  if(f.dc)badges+=`<span class="badge badge-deep">DEEP INDEXED</span> `;
  if(['py','sql','js','ts','jsx','sh'].includes(f.t))badges+=`<span class="badge badge-code">CODE</span> `;
  // Description line
  let descHtml=f.desc?`<div class="card-desc">${esc(f.desc)}</div>`:'';
  // Open button label
  const appName=APP_MAP[f.t]||'Default';
  const appLabel=appName==='Default'?'OPEN':appName.split(' ').pop();
  // Drive link
  const isGDrive=f.p.toLowerCase().includes('googledrive')||f.p.toLowerCase().includes('cloudstorage');
  const fname=encodeURIComponent(f.n.replace(/\\.[^.]+$/,''));
  const driveHtml=isGDrive?`<a class="btn btn-drive" href="https://drive.google.com/drive/search?q=${fname}" target="_blank">DRIVE</a>`:'';

  return `<div class="card">
    <div class="card-header"><div class="file-name">${name}</div><span class="type-badge" style="background:${tc}22;color:${tc};border:1px solid ${tc}44">.${esc(f.t)}</span>${f.d?`<span class="date-badge">${f.d}</span>`:''}</div>
    ${descHtml}
    <div class="card-meta"><span class="meta-item">${esc(f.c)}</span><span class="meta-item">${sz}</span>${badges}</div>
    ${ecoHtml}${topicHtml}${deepHtml}
    <div class="actions">
      <span class="btn btn-open" onclick="openFile('${esc(f.p)}','${esc(f.t)}')" title="Open in ${appName}">&#9654; ${appLabel}</span>
      <span class="btn btn-reveal" onclick="revealFile('${esc(f.p)}')" title="Reveal in Finder">REVEAL</span>
      <span class="btn btn-copy" onclick="copyPath('${esc(f.p)}',this)">COPY</span>
      ${driveHtml}
    </div>
    <div style="font-size:9px;color:#484f58;margin-top:5px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap" title="${esc(f.p)}">${esc(f.p)}</div>
  </div>`;
}
function esc(s){return String(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;').replace(/'/g,'&#39;');}
function escRe(s){return s.replace(/[.*+?^${}()|[\\]\\\\]/g,'\\\\$&');}

const BRIDGE='http://localhost:8181';let _bridgeOk=null;
async function checkBridge(){try{const r=await fetch(BRIDGE+'/health',{signal:AbortSignal.timeout(1500)});_bridgeOk=r.ok;}catch(e){_bridgeOk=false;}return _bridgeOk;}
async function bridgeCall(endpoint,params){if(_bridgeOk===null)await checkBridge();if(!_bridgeOk)return false;try{const qs=Object.entries(params).map(([k,v])=>k+'='+encodeURIComponent(v)).join('&');const r=await fetch(`${BRIDGE}/${endpoint}?${qs}`,{signal:AbortSignal.timeout(3000)});return r.ok;}catch(e){return false;}}

async function openFile(path,fileType){
  const app=APP_MAP[fileType]||'';
  const bridged=await bridgeCall('open',app?{path,app}:{path});
  if(bridged)return;
  const fileUrl='file://'+path;
  if(location.protocol==='file:'){window.open(fileUrl,'_blank');return;}
  showPathModal('Bridge offline. Open manually:',path);
}
// === BUG FIX (2026-05-03): REVEAL no longer dumps into folder of thousands ===
async function revealFile(path){
  // Bridge does `open -R "$path"` -> Finder opens parent AND highlights the file.
  const bridged=await bridgeCall('reveal',{path});
  if(bridged)return;
  // Bridge offline: ALWAYS surface the modal with the path so the user can use
  // Finder's Go > Go to Folder (Cmd+Shift+G). Never window.open the parent dir
  // -- that's what was dumping users into folders of thousands.
  showPathModal('Bridge offline. To reveal in Finder: copy the path below, then in Finder press Cmd+Shift+G and paste.',path);
}
// === END FIX ===
// === BUG FIX (2026-05-03): COPY now actually works on file:// URLs ===
async function copyPath(path,el){
  const done=()=>{el.textContent='COPIED';el.classList.add('copied');setTimeout(()=>{el.textContent='COPY';el.classList.remove('copied');},2000);};
  // Tier 1: bridge (works in any context)
  try{const bridged=await bridgeCall('copy',{text:path});if(bridged){done();return;}}catch(e){}
  // Tier 2: navigator.clipboard (works in https/localhost contexts)
  if(navigator.clipboard&&navigator.clipboard.writeText&&window.isSecureContext){
    try{await navigator.clipboard.writeText(path);done();return;}catch(e){}
  }
  // Tier 3: ALWAYS works -- visible textarea (offscreen textareas are blocked on file://)
  // + execCommand('copy'). If that also fails, surface the modal so the user can Cmd+C manually.
  const ta=document.createElement('textarea');
  ta.value=path; ta.style.cssText='position:fixed;top:0;left:0;width:1px;height:1px;opacity:0.01;z-index:9999';
  document.body.appendChild(ta); ta.focus(); ta.select();
  try{
    const ok=document.execCommand('copy');
    document.body.removeChild(ta);
    if(ok){done();return;}
  }catch(e){try{document.body.removeChild(ta);}catch(_){}}
  showPathModal('Auto-copy unavailable. Path is selected -- press Cmd+C to copy:',path);
}
// === END FIX ===
function showPathModal(msg,path){const ex=document.getElementById('path-modal');if(ex)ex.remove();const m=document.createElement('div');m.id='path-modal';m.style.cssText='position:fixed;top:0;left:0;right:0;bottom:0;background:#000a;z-index:9999;display:flex;align-items:center;justify-content:center';m.innerHTML=`<div style="background:#161b22;border:1px solid #388bfd;border-radius:10px;padding:24px;max-width:700px;width:90%"><div style="font-size:12px;color:#8b949e;margin-bottom:10px">${msg}</div><input style="width:100%;background:#0d1117;border:1px solid #30363d;color:#e6edf3;padding:8px;border-radius:5px;font-family:monospace" value="${path}" readonly onclick="this.select()"><button style="margin-top:12px;padding:6px 16px;background:#21262d;border:1px solid #30363d;color:#e6edf3;border-radius:5px;cursor:pointer" onclick="this.closest('#path-modal').remove()">Close</button></div>`;m.onclick=e=>{if(e.target===m)m.remove();};document.body.appendChild(m);}

function addTopicSearch(t){document.getElementById('search-box').value=t;state.search=t;applyFilters();}
function filterByEco(ecoId){state.ecos.clear();state.ecos.add(ecoId);document.querySelectorAll('.eco-chip').forEach(c=>c.classList.remove('active'));const chip=document.querySelector(`[data-eco="${ecoId}"]`);if(chip)chip.classList.add('active');applyFilters();}
function togglePill(el,group){const val=el.dataset.val;if(val==='all-'+group){document.querySelectorAll(`.pill`).forEach(p=>{if(p.closest('#'+group+'-pills'))p.classList.remove('active');});el.classList.add('active');if(group==='cat')state.cats.clear();else state.types.clear();}else{document.querySelector(`[data-val="all-${group}"]`).classList.remove('active');el.classList.toggle('active');const target=group==='cat'?state.cats:state.types;if(el.classList.contains('active'))target.add(val);else target.delete(val);if(target.size===0)document.querySelector(`[data-val="all-${group}"]`).classList.add('active');}applyFilters();}
function toggleEco(id){const chip=document.querySelector(`[data-eco="${id}"]`);if(state.ecos.has(id)){state.ecos.delete(id);chip.classList.remove('active');}else{state.ecos.add(id);chip.classList.add('active');}applyFilters();}
function toggleFilter(key){state[key]=!state[key];document.getElementById('toggle-'+key).classList.toggle('active',state[key]);applyFilters();}
function toggleSection(label){label.classList.toggle('collapsed');label.nextElementSibling.classList.toggle('hidden');}
function clearAllFilters(){state.search='';state.cats.clear();state.types.clear();state.dateFrom='';state.dateTo='';state.ecos.clear();state.session=false;state.ecoref=false;state.deep=false;state.code=false;document.getElementById('search-box').value='';document.getElementById('date-from').value='';document.getElementById('date-to').value='';document.querySelectorAll('.pill').forEach(p=>p.classList.remove('active'));document.querySelectorAll('[data-val="all-cat"],[data-val="all-type"]').forEach(p=>p.classList.add('active'));document.querySelectorAll('.eco-chip').forEach(c=>c.classList.remove('active'));document.querySelectorAll('.toggle-btn').forEach(b=>b.classList.remove('active'));applyFilters();}

async function checkRelay(){try{const r=await fetch('http://37.27.169.232:8080/health',{signal:AbortSignal.timeout(4000)});const b=document.getElementById('relay-badge');if(r.ok){b.textContent='RELAY ONLINE';b.classList.add('ok');}else b.textContent='RELAY ERR';}catch(e){document.getElementById('relay-badge').textContent='RELAY OFFLINE';}}

document.getElementById('search-box').addEventListener('input',e=>{state.search=e.target.value;applyFilters();});
document.getElementById('date-from').addEventListener('change',e=>{state.dateFrom=e.target.value;applyFilters();});
document.getElementById('date-to').addEventListener('change',e=>{state.dateTo=e.target.value;applyFilters();});
buildSidebar();applyFilters();
async function checkBridgeBadge(){const ok=await checkBridge();const b=document.getElementById('bridge-badge');if(ok){b.textContent='BRIDGE ONLINE';b.classList.add('ok');}else{b.textContent='BRIDGE OFFLINE';b.style.color='#f85149';}}
checkRelay();checkBridgeBadge();setInterval(checkRelay,30000);setInterval(checkBridgeBadge,15000);
</script>
</body>
</html>''')

print(f"\\nV8 written: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
print(f"Total files: {len(FILES)}")
print(f"Deep-indexed: {enhanced}")
print(f"New ecosystem assignments: {new_eco_assignments}")
