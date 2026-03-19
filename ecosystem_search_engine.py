#!/usr/bin/env python3
"""
BroyhillGOP Advanced Ecosystem Search Engine v2
================================================
Searches ALL file types: .md, .docx, .sql, .py, .html, .json,
.xlsx, .numbers, .csv, .txt, .sh

Extracts ecosystem references, component features, topics.
Auto-updates MASTER_DIRECTORY_INDEX.md with new files found.
Outputs structured JSON + gap analysis MD + V5 HTML data refresh.

v2 Changes (Mar 18, 2026):
- Scans .docx via python-docx (text extraction)
- Scans .sql, .py, .html, .json, .sh, .txt as text
- Scans .xlsx via openpyxl (sheet names + cell text)
- Scans .csv (header row + sample)
- Auto-appends new files to MASTER_DIRECTORY_INDEX.md
- Topic extraction from file content
"""

import os
import re
import json
import sys
from collections import defaultdict
from pathlib import Path
from datetime import datetime

# Optional imports for binary formats
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. DOCX files will be skipped.")
    print("  Install: pip3 install python-docx")

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False
    print("WARNING: openpyxl not installed. XLSX files will be skipped.")
    print("  Install: pip3 install openpyxl")

# ============================================================
# CONFIGURATION
# ============================================================

SEARCH_DIRS = [
    "/Users/Broyhill/Desktop/BroyhillGOP-CURSOR",
    "/Users/Broyhill/Desktop/NCBOE-FEC- DONORS 2015-2026",
    "/Users/Broyhill/Downloads",
    "/Users/Broyhill/Documents",
]

OUTPUT_DIR = "/Users/Broyhill/Desktop/BroyhillGOP-CURSOR"
MASTER_INDEX_PATH = "/Users/Broyhill/Desktop/BroyhillGOP-CURSOR/MASTER_DIRECTORY_INDEX.md"

# File extensions to scan
TEXT_EXTENSIONS = {'.md', '.sql', '.py', '.html', '.json', '.txt', '.sh', '.css', '.js', '.jsx', '.ts'}
BINARY_EXTENSIONS = {'.docx', '.xlsx', '.csv', '.numbers'}
ALL_EXTENSIONS = TEXT_EXTENSIONS | BINARY_EXTENSIONS

# Directories to skip
SKIP_DIRS = {'node_modules', '.git', '.vercel', '__pycache__', '.next', '.cursor', 'venv'}

# Known ecosystem mapping (E00-E57 + E58)
ECOSYSTEM_NAMES = {
    "E00": "DataHub / Master Database",
    "E01": "Donor Intelligence Engine",
    "E02": "Donation Processing",
    "E03": "Candidate Profiles / Marketing Automation",
    "E04": "Activist Network",
    "E05": "Volunteer Management",
    "E06": "Analytics Engine",
    "E07": "Issue Tracking",
    "E08": "Communications Library",
    "E09": "Content Creation AI",
    "E10": "Compliance Manager",
    "E11": "Budget Management / Training LMS",
    "E12": "Campaign Operations",
    "E13": "AI Hub",
    "E14": "Print Production",
    "E15": "Contact Directory",
    "E16": "TV/Radio AI + Voice Synthesis",
    "E17": "RVM (Ringless Voicemail)",
    "E18": "VDP / Print Advertising",
    "E19": "Social Media Manager / Personalization",
    "E20": "Intelligence Brain",
    "E21": "ML Clustering",
    "E22": "A/B Testing Engine",
    "E23": "Creative Asset / 3D Engine",
    "E24": "Candidate Portal",
    "E25": "Donor Portal",
    "E26": "Volunteer Portal",
    "E27": "Realtime Dashboard",
    "E28": "Financial Dashboard",
    "E29": "Analytics Dashboard",
    "E30": "Email Engine",
    "E31": "SMS/MMS Engine",
    "E32": "Phone Banking",
    "E33": "Direct Mail",
    "E34": "Events",
    "E35": "Interactive Comm Hub",
    "E36": "Messenger Integration",
    "E37": "Event Management",
    "E38": "Volunteer Coordination",
    "E39": "P2P Fundraising",
    "E40": "Automation Control Panel",
    "E41": "Campaign Builder",
    "E42": "News Intelligence",
    "E43": "Advocacy Tools",
    "E44": "Vendor Compliance / Security",
    "E45": "Video Studio",
    "E46": "Broadcast Hub",
    "E47": "AI Script Generator / Unified Voice",
    "E48": "Communication DNA",
    "E49": "Interview System",
    "E50": "GPU Orchestrator",
    "E51": "Nexus Dashboard",
    "E52": "Contact Intelligence Engine",
    "E53": "Document Generation",
    "E54": "Calendar/Scheduling",
    "E55": "API Gateway",
    "E56": "Visitor Deanonymization",
    "E57": "Messaging Center",
    "E58": "Business Model / Candidate Network",
}

# Topic keywords to extract from content
TOPIC_KEYWORDS = [
    "donor", "donation", "FEC", "compliance", "volunteer", "candidate",
    "campaign", "analytics", "dashboard", "email", "SMS", "phone",
    "direct mail", "events", "fundraising", "voter", "DataTrust",
    "identity resolution", "dedup", "matching", "spine", "contribution",
    "committee", "NC BOE", "NCGOP", "WinRed", "scoring", "grading",
    "segmentation", "archetype", "microsegment", "cultivation",
    "enrichment", "geocodio", "USPS", "address", "Nexus", "brain",
    "AI", "ML", "clustering", "voice", "video", "print", "social media",
    "automation", "workflow", "API", "webhook", "integration",
    "RunPod", "GPU", "Supabase", "PostgreSQL", "migration",
    "heat map", "precinct", "district", "county", "municipal",
    "sheriff", "mayor", "council", "BOE", "school board",
]

# ============================================================
# FILE CONTENT EXTRACTORS
# ============================================================

def extract_text_file(filepath):
    """Read a text-based file."""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        return None

def extract_docx(filepath):
    """Extract text from a .docx file."""
    if not HAS_DOCX:
        return None
    try:
        doc = DocxDocument(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also get table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return '\n'.join(paragraphs)
    except Exception as e:
        return None

def extract_xlsx(filepath):
    """Extract text from .xlsx headers and first rows."""
    if not HAS_XLSX:
        return None
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}")
            row_count = 0
            for row in ws.iter_rows(max_row=20, values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    text_parts.append(' | '.join(cells))
                row_count += 1
        wb.close()
        return '\n'.join(text_parts)
    except Exception as e:
        return None

def extract_csv_header(filepath):
    """Extract CSV header row and first few data rows."""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            lines = []
            for i, line in enumerate(f):
                if i > 10:
                    break
                lines.append(line.strip())
            return '\n'.join(lines)
    except Exception as e:
        return None

def extract_content(filepath):
    """Route to the right extractor based on extension."""
    ext = Path(filepath).suffix.lower()
    if ext in TEXT_EXTENSIONS:
        return extract_text_file(filepath)
    elif ext == '.docx':
        return extract_docx(filepath)
    elif ext == '.xlsx' or ext == '.xls':
        return extract_xlsx(filepath)
    elif ext == '.csv':
        return extract_csv_header(filepath)
    elif ext == '.numbers':
        return f"Apple Numbers file: {Path(filepath).name}"
    return None

def extract_topics(content, filename):
    """Extract topic keywords found in content."""
    if not content:
        return []
    content_lower = content.lower()
    fname_lower = filename.lower()
    found = []
    for kw in TOPIC_KEYWORDS:
        if kw.lower() in content_lower or kw.lower() in fname_lower:
            found.append(kw)
    return found[:15]  # Cap at 15 topics

# ============================================================
# SEARCH ENGINE
# ============================================================

class EcosystemSearchEngine:
    def __init__(self):
        self.results = {
            "ecosystem_refs": defaultdict(lambda: defaultdict(list)),
            "files_scanned": 0,
            "files_with_hits": 0,
            "total_ecosystem_hits": 0,
            "errors": [],
            "file_index": [],  # NEW: full file inventory with topics
        }

    def is_noise(self, line, match_pos):
        """Filter out URL/footnote noise."""
        url_pattern = r'https?://[^\s\)]*'
        for m in re.finditer(url_pattern, line):
            if m.start() <= match_pos <= m.end():
                return True
        return False

    def extract_context(self, lines, line_idx, window=2):
        """Extract context window around a match."""
        start = max(0, line_idx - window)
        end = min(len(lines), line_idx + window + 1)
        ctx = "\n".join(lines[start:end]).strip()
        return ctx[:400] + "..." if len(ctx) > 400 else ctx

    def scan_file(self, filepath):
        """Scan a single file for ecosystem refs and index it."""
        content = extract_content(filepath)
        if content is None:
            return

        self.results["files_scanned"] += 1
        file_had_hits = False
        fname = Path(filepath).name
        ext = Path(filepath).suffix.lower().lstrip('.')

        # Extract topics
        topics = extract_topics(content, fname)

        # Add to file index
        self.results["file_index"].append({
            "name": fname,
            "path": str(filepath),
            "type": ext,
            "topics": ", ".join(topics),
            "size": os.path.getsize(filepath),
        })

        # Ecosystem ID search
        lines = content.split('\n')
        eco_pattern = r'\bE(\d{2})\b'
        for i, line in enumerate(lines):
            for m in re.finditer(eco_pattern, line):
                eco_num = m.group(1)
                eco_id = f"E{eco_num}"
                if eco_id not in ECOSYSTEM_NAMES:
                    continue
                if self.is_noise(line, m.start()):
                    continue
                ctx = self.extract_context(lines, i)
                self.results["ecosystem_refs"][eco_id][str(filepath)].append({
                    "line": i + 1,
                    "match": line.strip()[:200],
                })
                self.results["total_ecosystem_hits"] += 1
                file_had_hits = True

        if file_had_hits:
            self.results["files_with_hits"] += 1

    def scan_directory(self, dirpath):
        """Recursively scan a directory for ALL supported file types."""
        dirpath = Path(dirpath)
        if not dirpath.exists():
            self.results["errors"].append(f"Directory not found: {dirpath}")
            return

        count_by_ext = defaultdict(int)
        for root, dirs, files in os.walk(dirpath):
            # Skip unwanted directories
            dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
            for fname in files:
                ext = Path(fname).suffix.lower()
                if ext in ALL_EXTENSIONS:
                    count_by_ext[ext] += 1
                    self.scan_file(os.path.join(root, fname))

        ext_summary = ", ".join(f"{ext}: {c}" for ext, c in sorted(count_by_ext.items()))
        print(f"  Scanned in {dirpath.name}: {ext_summary}")

    def run(self):
        """Run the full search across all configured directories."""
        print("=" * 70)
        print("BroyhillGOP Ecosystem Search Engine v2")
        print(f"Scanning: {', '.join(e for e in sorted(ALL_EXTENSIONS))}")
        print("=" * 70)
        for d in SEARCH_DIRS:
            print(f"\nScanning: {d}")
            self.scan_directory(d)
        print(f"\n{'=' * 70}")
        print(f"SCAN COMPLETE")
        print(f"  Files scanned: {self.results['files_scanned']}")
        print(f"  Files with ecosystem refs: {self.results['files_with_hits']}")
        print(f"  Ecosystem hits: {self.results['total_ecosystem_hits']}")
        print(f"  Unique files indexed: {len(self.results['file_index'])}")
        print(f"  Errors: {len(self.results['errors'])}")
        print(f"{'=' * 70}")

    def generate_cross_reference(self):
        """Build planned-vs-built cross-reference."""
        cursor_dir = Path("/Users/Broyhill/Desktop/BroyhillGOP-CURSOR/ecosystems")
        backend_dir = Path("/Users/Broyhill/Desktop/BroyhillGOP-CURSOR/backend/python/ecosystems")
        built_files = {}
        for search_dir in [cursor_dir, backend_dir]:
            if search_dir.exists():
                for f in search_dir.iterdir():
                    if f.name.startswith("ecosystem_") and f.suffix == ".py":
                        m = re.search(r'ecosystem_(\d+)', f.name)
                        if m:
                            eco_id = f"E{int(m.group(1)):02d}"
                            if eco_id not in built_files:
                                built_files[eco_id] = []
                            built_files[eco_id].append(f.name)

        cross_ref = {}
        for eco_id, eco_name in sorted(ECOSYSTEM_NAMES.items()):
            num_refs = sum(len(ctxs) for ctxs in self.results["ecosystem_refs"].get(eco_id, {}).values())
            num_files = len(self.results["ecosystem_refs"].get(eco_id, {}))
            has_code = eco_id in built_files
            cross_ref[eco_id] = {
                "name": eco_name,
                "has_code": has_code,
                "code_files": built_files.get(eco_id, []),
                "doc_references": num_refs,
                "files_mentioning": num_files,
                "status": "BUILT" if has_code else ("PLANNED" if num_refs > 0 else "UNKNOWN"),
            }
        return cross_ref

    def write_json_output(self, cross_ref):
        """Write structured JSON output with full file index."""
        output = {
            "generated": datetime.now().isoformat(),
            "scan_summary": {
                "files_scanned": self.results["files_scanned"],
                "files_with_ecosystem_refs": self.results["files_with_hits"],
                "total_ecosystem_hits": self.results["total_ecosystem_hits"],
                "unique_files_indexed": len(self.results["file_index"]),
                "errors_count": len(self.results["errors"]),
                "extensions_scanned": sorted(list(ALL_EXTENSIONS)),
            },
            "file_index": self.results["file_index"],
            "ecosystem_cross_reference": cross_ref,
            "errors": self.results["errors"][:20],
        }
        outpath = os.path.join(OUTPUT_DIR, "ecosystem_search_results.json")
        with open(outpath, 'w') as f:
            json.dump(output, f, indent=2, default=str)
        print(f"\nJSON output: {outpath}")
        return outpath

    def write_gap_analysis(self, cross_ref):
        """Write human-readable gap analysis MD."""
        outpath = os.path.join(OUTPUT_DIR, "ECOSYSTEM_GAP_ANALYSIS.md")
        lines = []
        lines.append("# BroyhillGOP Ecosystem Gap Analysis v2")
        lines.append(f"## Generated {datetime.now().strftime('%B %d, %Y')}\n")
        lines.append(f"**Files scanned:** {self.results['files_scanned']}")
        lines.append(f"**File types:** {', '.join(sorted(ALL_EXTENSIONS))}")
        lines.append(f"**Files with ecosystem refs:** {self.results['files_with_hits']}")
        lines.append(f"**Total ecosystem hits:** {self.results['total_ecosystem_hits']}")
        lines.append(f"**Unique files indexed:** {len(self.results['file_index'])}\n")
        lines.append("---\n")

        # Cross-reference table
        lines.append("## Ecosystem Cross-Reference: Planned vs Built\n")
        lines.append("| ID | Name | Status | Code Files | Doc Refs |")
        lines.append("|-----|------|--------|------------|----------|")
        for eco_id in sorted(cross_ref.keys(), key=lambda x: int(x[1:])):
            r = cross_ref[eco_id]
            icon = "BUILT" if r["status"] == "BUILT" else ("PLANNED" if r["status"] == "PLANNED" else "UNKNOWN")
            lines.append(f"| {eco_id} | {r['name']} | {icon} | {len(r['code_files'])} | {r['doc_references']} |")

        # File type breakdown
        lines.append("\n---\n")
        lines.append("## File Type Breakdown\n")
        type_counts = defaultdict(int)
        for f in self.results["file_index"]:
            type_counts[f["type"]] += 1
        for ext, count in sorted(type_counts.items(), key=lambda x: -x[1]):
            lines.append(f"- **.{ext}**: {count} files")

        with open(outpath, 'w') as f:
            f.write("\n".join(lines))
        print(f"Gap analysis: {outpath}")

    def update_master_index(self):
        """Append newly discovered files to MASTER_DIRECTORY_INDEX.md."""
        if not os.path.exists(MASTER_INDEX_PATH):
            print("MASTER_DIRECTORY_INDEX.md not found, skipping update.")
            return

        # Read existing index to find already-listed files
        with open(MASTER_INDEX_PATH, 'r') as f:
            existing_content = f.read()

        existing_names = set()
        for line in existing_content.split('\n'):
            # Extract filenames from markdown table rows and text
            matches = re.findall(r'[\w.-]+\.(?:py|sql|md|docx|html|json|csv|xlsx|sh)', line)
            existing_names.update(matches)

        # Find new files not in the index
        new_files = []
        for fi in self.results["file_index"]:
            if fi["name"] not in existing_names and fi["topics"]:
                new_files.append(fi)

        if not new_files:
            print(f"Master index is up to date ({len(existing_names)} files tracked).")
            return

        # Append new files section
        with open(MASTER_INDEX_PATH, 'a') as f:
            f.write(f"\n\n---\n\n## AUTO-DISCOVERED FILES ({datetime.now().strftime('%B %d, %Y')})\n\n")
            f.write(f"Search engine v2 found {len(new_files)} files not in the index:\n\n")
            f.write("| File | Type | Topics |\n")
            f.write("|------|------|--------|\n")
            for fi in sorted(new_files, key=lambda x: x["type"]):
                f.write(f"| {fi['name']} | .{fi['type']} | {fi['topics'][:80]} |\n")

        print(f"Appended {len(new_files)} new files to MASTER_DIRECTORY_INDEX.md")

    def generate_v5_data(self):
        """Generate a JS data file that V5 HTML can load to auto-refresh."""
        outpath = os.path.join(OUTPUT_DIR, "god_file_v5_data.js")
        entries = []
        for fi in self.results["file_index"]:
            # Determine category from path
            path = fi["path"]
            if "/ECOSYSTEM_REPORTS/" in path:
                cat = "Ecosystem Reports"
            elif "/migrations/" in path:
                cat = "SQL Migrations"
            elif "/backend/python/ecosystems/" in path or "/ecosystems/" in path:
                cat = "Python Ecosystems"
            elif "/backend/python/integrations/" in path:
                cat = "Python Integrations"
            elif "/backend/python/engines/" in path:
                cat = "Python Engines"
            elif "/pipeline/" in path:
                cat = "Pipeline Scripts"
            elif "/scripts/" in path:
                cat = "Operational Scripts"
            elif "/docs/" in path:
                cat = "Documentation"
            elif "SESSION-STATE" in fi["name"]:
                cat = "Session Handoffs"
            elif fi["type"] in ("csv", "xlsx", "numbers"):
                cat = "Raw Data"
            else:
                cat = "Other"

            entries.append({
                "name": fi["name"],
                "path": fi["path"],
                "cat": cat,
                "type": fi["type"],
                "topics": fi["topics"],
            })

        js_content = f"// Auto-generated {datetime.now().isoformat()}\nconst FILES_AUTO = {json.dumps(entries, indent=2)};\n"
        with open(outpath, 'w') as f:
            f.write(js_content)
        print(f"V5 data file: {outpath} ({len(entries)} entries)")


# ============================================================
# MAIN
# ============================================================

def main():
    engine = EcosystemSearchEngine()
    engine.run()
    cross_ref = engine.generate_cross_reference()
    engine.write_json_output(cross_ref)
    engine.write_gap_analysis(cross_ref)
    engine.update_master_index()
    engine.generate_v5_data()

    # Quick summary
    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    built = [k for k, v in cross_ref.items() if v["status"] == "BUILT"]
    planned = [k for k, v in cross_ref.items() if v["status"] == "PLANNED"]
    print(f"  BUILT: {len(built)} ecosystems")
    print(f"  PLANNED only: {len(planned)} ecosystems")
    print(f"  Files indexed: {len(engine.results['file_index'])}")

    type_counts = defaultdict(int)
    for f in engine.results["file_index"]:
        type_counts[f["type"]] += 1
    for ext, count in sorted(type_counts.items(), key=lambda x: -x[1]):
        print(f"    .{ext}: {count}")

    print(f"\nOutputs:")
    print(f"  ecosystem_search_results.json")
    print(f"  ECOSYSTEM_GAP_ANALYSIS.md")
    print(f"  MASTER_DIRECTORY_INDEX.md (auto-updated)")
    print(f"  god_file_v5_data.js (for V5 HTML)")


if __name__ == "__main__":
    main()
